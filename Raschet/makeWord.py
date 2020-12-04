# -*- coding: utf-8 -*-

import CalcClass
import docx
import decimal
import os


def makeWord_obvn(data_in:CalcClass.data_in, data_out:CalcClass.data_out, docum=None):
    if os.path.isfile(docum):
        doc = docx.Document(docum)
    else:
        doc = docx.Document('temp.docx')
    doc.add_heading(f'Расчет на прочность обечайки {data_in.name}, нагруженной внутренним избыточным давлением').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture('pic/ObCil.gif')

    doc.add_paragraph('Исходные данные').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=0, cols=0)
    table.add_column(7000000)
    table.add_column(750000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.add_column(750000)
    table.add_row()  
    table.cell(0, 0).text = 'Наименование и размерность'
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = 'Обозначение'
    table.cell(0, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Значение'
    table.add_row()
    table.cell(1, 0).text = 'Расчетная температура, °С'
    table.cell(1, 1).text = 't'
    table.cell(1, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 2).text = f'{data_in.temp}'
    table.add_row()
    table.cell(2, 0).text = 'Расчетное внутреннее давление, МПа'
    table.cell(2, 1).text ='p'
    table.cell(2, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 2).text =f'{data_in.press}'
    table.add_row()
    table.cell(3, 0).text = 'Марка стали'
    table.cell(3, 1).text =''
    table.cell(3, 2).text =f'{data_in.steel}'
    table.add_row()
    table.cell(4, 0).text = 'Допускаемое напряжение при расчетной температуре, МПа'
    table.cell(4, 1).text = '[σ]'
    table.cell(4, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4, 2).text =f'{data_in.sigma_d}'
    table.add_row()
    table.cell(5, 0).text = 'Коэффициент прочности сварного шва'
    table.cell(5, 1).paragraphs[0].add_run('φ_p').font.math = True
    #table.cell(5, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(5, 2).text =f'{data_in.fi}'
    table.add_row()
    table.cell(6, 0).text = 'Внутренний диаметр обечайки, мм'
    table.cell(6, 1).text = 'D'
    table.cell(6, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(6, 2).text =f'{data_in.dia}'
    table.add_row()
    table.cell(7, 0).text = 'Прибавка на коррозию, мм'
    table.cell(7, 1).paragraphs[0].add_run('c_1').font.math = True
    #table.cell(7, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(7, 2).text =f'{data_in.c_kor}'
    table.add_row()
    table.cell(8, 0).text = 'Прибавка для компенсации минусового допуска листа, мм'
    table.cell(8, 1).paragraphs[0].add_run('c_2').font.math = True
    #table.cell(8, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(8, 2).text =f'{data_in.c_minus}'
    table.add_row()
    table.cell(9, 0).text = 'Технологическая прибавка, мм'
    table.cell(9, 1).paragraphs[0].add_run('c_3').font.math = True
    #table.cell(9, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(9, 2).text =f'{data_in.c_3}'
    doc.add_paragraph('')
    doc.add_paragraph('Результаты расчета').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    doc.add_paragraph('Толщину стенки вычисляют по формуле:')
    doc.add_paragraph().add_run('s≥s_p+c').font.math = True
    p = doc.add_paragraph('где ')
    p.add_run('s_p').font.math = True
    p.add_run(' - расчетная толщина стенки обечайки')
    doc.add_paragraph().add_run('s_p=(p∙D)/(2∙[σ]∙φ_p-p)').font.math = True
    doc.add_paragraph('c - сумма прибавок к расчетной толщине')
    doc.add_paragraph().add_run('c=c_1+c_2+c_3').font.math = True
    doc.add_paragraph().add_run(f'c={data_in.c_kor}+{data_in.c_minus}+{data_in.c_3}={data_out.c:.2f} мм').font.math = True
    doc.add_paragraph().add_run(f's_p=({data_in.press}∙{data_in.dia})/(2∙{data_in.sigma_d}∙{data_in.fi}-{data_in.press})={data_out.s_calcr:.2f} мм').font.math = True
    doc.add_paragraph().add_run(f's={data_out.s_calcr:.2f}+{data_out.c:.2f}={data_out.s_calc:.2f} мм').font.math = True
    if data_in.s_prin > data_out.s_calc:
        doc.add_paragraph(f'Принятая толщина s={data_in.s_prin} мм')
    else:
        doc.add_paragraph().add_run(f'Принятая толщина s={data_in.s_prin} мм').font.color.rgb = docx.shared.RGBColor(255,0,0)
    doc.add_paragraph('Допускаемое внутреннее избыточное давление вычисляют по формуле:')
    doc.add_paragraph().add_run('[p]=(2∙[σ]∙φ_p∙(s-c))/(D+s-c)').font.math = True
    doc.add_paragraph().add_run(f'[p]=(2∙{data_in.sigma_d}∙{data_in.fi}∙({data_in.s_prin}-{data_out.c:.2f}))/({data_in.dia}+{data_in.s_prin}-{data_out.c:.2f})={data_out.press_d:.2f} МПа').font.math = True
    doc.add_paragraph().add_run('[p]≥p').font.math = True
    doc.add_paragraph().add_run(f'{data_out.press_d:.2f}≥{data_in.press}').font.math = True
    if data_out.press_d > data_in.press:
        doc.add_paragraph('Условие прочности выполняется')
    else:
        doc.add_paragraph().add_run('Условие прочности не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)
    p = doc.add_paragraph('Границы применения формул ')
    if data_in.dia >= 200:
        p.add_run('при D ≥ 200 мм')
        doc.add_paragraph().add_run('(s-c)/(D)≤0.1').font.math = True
        doc.add_paragraph().add_run(f'({data_in.s_prin}-{data_out.c:.2f})/({data_in.dia})={(data_in.s_prin-data_out.c)/data_in.dia:.3f}≤0.1').font.math = True
    else:
        p.add_run('при D < 200 мм')
        doc.add_paragraph().add_run('(s-c)/(D)≤0.3').font.math = True
        doc.add_paragraph().add_run(f'({data_in.s_prin}-{data_out.c:.2f})/({data_in.dia})={(data_in.s_prin-data_out.c)/data_in.dia:.2f}≤0.3').font.math = True

    doc.save(docum)

def makeWord_obnar(data_in:CalcClass.data_in, data_out:CalcClass.data_out, docum=None):
    doc = docx.Document(docum)
    doc.add_heading(f'Расчет на прочность обечайки {data_in.name}, нагруженной наружным давлением').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture('pic/ObCil.gif')
    doc.add_paragraph('Исходные данные').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=0, cols=0)
    table.add_column(7000000)
    table.add_column(750000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.add_column(750000)
    table.add_row()  
    table.cell(0, 0).text = 'Наименование и размерность'
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = 'Обозначение'
    table.cell(0, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Значение'
    table.add_row()
    table.cell(1, 0).text = 'Расчетная температура, °С'
    table.cell(1, 1).text = 't'
    table.cell(1, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 2).text = f'{data_in.temp}'
    table.add_row()
    table.cell(2, 0).text = 'Расчетное внутреннее давление, МПа'
    table.cell(2, 1).text ='p'
    table.cell(2, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 2).text =f'{data_in.press}'
    table.add_row()
    table.cell(3, 0).text = 'Марка стали'
    table.cell(3, 1).text =''
    table.cell(3, 2).text =f'{data_in.steel}'
    table.add_row()
    table.cell(4, 0).text = 'Допускаемое напряжение при расчетной температуре, МПа'
    table.cell(4, 1).text = '[σ]'
    table.cell(4, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4, 2).text =f'{data_in.sigma_d}'
    table.add_row()
    table.cell(5, 0).text = 'Модуль продольной упругости при расчетной темп-ре, МПа'
    table.cell(5, 1).text = 'E'
    table.cell(5, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(5, 2).text =f'{data_in.E}'
    table.add_row()
    table.cell(6, 0).text = 'Коэффициент прочности сварного шва'
    table.cell(6, 1).paragraphs[0].add_run('φ_p').font.math = True
    #table.cell(6, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(6, 2).text =f'{data_in.fi}'
    table.add_row()
    table.cell(7, 0).text = 'Внутренний диаметр обечайки, мм'
    table.cell(7, 1).text = 'D'
    table.cell(7, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(7, 2).text =f'{data_in.dia}'
    table.add_row()
    table.cell(8, 0).text = 'Длина обечайки, мм'
    table.cell(8, 1).text = 'l'
    table.cell(8, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(8, 2).text =f'{data_in.l}'
    table.add_row()
    table.cell(9, 0).text = 'Прибавка на коррозию, мм'
    table.cell(9, 1).paragraphs[0].add_run('c_1').font.math = True
    #table.cell(9, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(9, 2).text =f'{data_in.c_kor}'
    table.add_row()
    table.cell(10, 0).text = 'Прибавка для компенсации минусового допуска листа, мм'
    table.cell(10, 1).paragraphs[0].add_run('c_2').font.math = True
    #table.cell(10, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(10, 2).text =f'{data_in.c_minus}'
    table.add_row()
    table.cell(11, 0).text = 'Технологическая прибавка, мм'
    table.cell(11, 1).paragraphs[0].add_run('c_3').font.math = True
    #table.cell(11, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(11, 2).text =f'{data_in.c_3}'
    doc.add_paragraph('')
    doc.add_paragraph('Результаты расчета').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    doc.add_paragraph('Толщину стенки вычисляют по формуле:')
    doc.add_paragraph().add_run('s≥s_p+c').font.math = True
    p = doc.add_paragraph('где ')
    p.add_run('s_p').font.math = True
    p.add_run(' - расчетная толщина стенки обечайки')
    doc.add_paragraph().add_run('s_p=max{1.06∙(10^-2∙D)/(B)∙(p/(10^-5∙E)∙l/D)^0.4;(1.2∙p∙D)/(2∙[σ]-p)}').font.math = True
    doc.add_paragraph('Коэффициент B вычисляют по формуле:')
    doc.add_paragraph().add_run('B=max{1;0.47∙(p/(10^-5∙E))^0.067∙(l/D)^0.4}').font.math = True
    doc.add_paragraph().add_run(f'0.47∙({data_in.press}/(10^-5∙{data_in.E}))^0.067∙({data_out.l}/{data_in.dia})^0.4={data_out.b_2:.2f}').font.math = True
    doc.add_paragraph().add_run(f'B=max(1;{data_out.b_2:.2f})={data_out.b:.2f}').font.math = True
    doc.add_paragraph('c - сумма прибавок к расчетной толщине')
    doc.add_paragraph().add_run('c=c_1+c_2+c_3').font.math = True
    doc.add_paragraph().add_run(f'c={data_in.c_kor}+{data_in.c_minus}+{data_in.c_3}={data_out.c:.2f} мм').font.math = True
    
    doc.add_paragraph().add_run(f'1.06∙(10^-2∙{data_in.dia})/({data_out.b:.2f})∙({data_in.press}/(10^-5∙{data_in.E})∙{data_out.l}/{data_in.dia})^0.4={data_out.s_calcr1:.2f}').font.math = True
    doc.add_paragraph().add_run(f'(1.2∙{data_in.press}∙{data_in.dia})/(2∙{data_in.sigma_d}-{data_in.press})={data_out.s_calcr2:.2f}').font.math = True
    doc.add_paragraph().add_run(f's_p=max({data_out.s_calcr1:.2f};{data_out.s_calcr2:.2f})={data_out.s_calcr:.2f} мм').font.math = True

    doc.add_paragraph().add_run(f's={data_out.s_calcr:.2f}+{data_out.c:.2f}={data_out.s_calc:.2f} мм').font.math = True
    if data_in.s_prin > data_out.s_calc:
        doc.add_paragraph(f'Принятая толщина s={data_in.s_prin} мм')
    else:
        doc.add_paragraph().add_run(f'Принятая толщина s={data_in.s_prin} мм').font.color.rgb = docx.shared.RGBColor(255,0,0)
    doc.add_paragraph('Допускаемое наружное давление вычисляют по формуле:')
    doc.add_paragraph().add_run('[p]=[p]_П/√(1+([p]_П/[p]_E)^2)').font.math = True
    doc.add_paragraph('допускаемое давление из условия прочности вычисляют по формуле:')
    doc.add_paragraph().add_run('[p]_П=(2∙[σ]∙(s-c))/(D+s-c)').font.math = True
    doc.add_paragraph().add_run(f'[p]_П=(2∙{data_in.sigma_d}∙({data_in.s_prin}-{data_out.c:.2f}))/({data_in.dia}+{data_in.s_prin}-{data_out.c:.2f})={data_out.press_dp:.2f} МПа').font.math = True
    doc.add_paragraph('допускаемое давление из условия устойчивости в пределах упругости вычисляют по формуле:')
    doc.add_paragraph().add_run('[p]_E=(2.08∙10^-5∙E)/(n_y∙B_1)∙D/l∙[(100∙(s-c))/D]^2.5').font.math = True
    p = doc.add_paragraph('коэффициент ')
    p.add_run('B_1').font.math = True
    p.add_run(' вычисляют по формуле')
    doc.add_paragraph().add_run('B_1=min{1;9.45∙D/l∙√(D/(100∙(s-c)))}').font.math = True
    doc.add_paragraph().add_run(f'9.45∙{data_in.dia}/{data_out.l}∙√({data_in.dia}/(100∙({data_in.s_prin}-{data_out.c:.2f})))={data_out.b1_2:.2f}').font.math = True
    doc.add_paragraph().add_run(f'B_1=min(1;{data_out.b1_2:.2f})={data_out.b1:.1f}').font.math = True
    doc.add_paragraph().add_run(f'[p]_E=(2.08∙10^-5∙{data_in.E})/({data_in.ny}∙{data_out.b1:.1f})∙{data_in.dia}/{data_out.l}∙[(100∙({data_in.s_prin}-{data_out.c:.2f}))/{data_in.dia}]^2.5={data_out.press_de:.2f} МПа').font.math = True
    doc.add_paragraph().add_run(f'[p]={data_out.press_dp:.2f}/√(1+({data_out.press_dp:.2f}/{data_out.press_de:.2f})^2)={data_out.press_d:.2f} МПа').font.math = True
    doc.add_paragraph().add_run('[p]≥p').font.math = True
    doc.add_paragraph().add_run(f'{data_out.press_d:.2f}≥{data_in.press}').font.math = True
    if data_out.press_d > data_in.press:
        doc.add_paragraph('Условие прочности выполняется')
    else:
        doc.add_paragraph().add_run('Условие прочности не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)
    p = doc.add_paragraph('Границы применения формул ')
    if data_in.dia >= 200:
        p.add_run('при D ≥ 200 мм')
        doc.add_paragraph().add_run('(s-c)/(D)≤0.1').font.math = True
        doc.add_paragraph().add_run(f'({data_in.s_prin}-{data_out.c:.2f})/({data_in.dia})={(data_in.s_prin-data_out.c)/data_in.dia:.3f}≤0.1').font.math = True
    else:
        p.add_run('при D < 200 мм')
        doc.add_paragraph().add_run('(s-c)/(D)≤0.3').font.math = True
        doc.add_paragraph().add_run(f'({data_in.s_prin}-{data_out.c:.2f})/({data_in.dia})={(data_in.s_prin-data_out.c)/data_in.dia:.2f}≤0.3').font.math = True

    doc.save(docum)

def makeWord_elvn(data_in:CalcClass.data_in, data_out:CalcClass.data_out, docum=None):

    doc = docx.Document(docum)
    # добавить полусферическое
    doc.add_heading(f'Расчет на прочность эллиптического днища {data_in.name}, нагруженного внутренним избыточным давлением').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Исходные данные').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=0, cols=0)
    table.add_column(7000000)
    table.add_column(750000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.add_column(750000)
    table.add_row()  
    table.cell(0, 0).text = 'Наименование и размерность'
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = 'Обозначение'
    table.cell(0, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Значение'
    table.add_row()
    table.cell(1, 0).text = 'Расчетная температура, °С'
    table.cell(1, 1).text = 't'
    table.cell(1, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 2).text = f'{data_in.temp}'
    table.add_row()
    table.cell(2, 0).text = 'Расчетное внутреннее давление, МПа'
    table.cell(2, 1).text ='p'
    table.cell(2, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 2).text =f'{data_in.press}'
    table.add_row()
    table.cell(3, 0).text = 'Марка стали'
    table.cell(3, 1).text =''
    table.cell(3, 2).text =f'{data_in.steel}'
    table.add_row()
    table.cell(4, 0).text = 'Допускаемое напряжение при расчетной температуре, МПа'
    table.cell(4, 1).text = '[σ]'
    table.cell(4, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4, 2).text =f'{data_in.sigma_d}'
    table.add_row()
    table.cell(5, 0).text = 'Коэффициент прочности сварного шва'
    table.cell(5, 1).paragraphs[0].add_run('φ_p').font.math = True
    #table.cell(5, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(5, 2).text =f'{data_in.fi}'
    table.add_row()
    table.cell(6, 0).text = 'Внутренний диаметр днища, мм'
    table.cell(6, 1).text = 'D'
    table.cell(6, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(6, 2).text =f'{data_in.dia}'
    table.add_row()
    table.cell(7, 0).text = 'Высота выпуклой части, мм'
    table.cell(7, 1).text = 'Н'
    table.cell(7, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(7, 2).text =f'{data_in.elH}'
    table.add_row()
    table.cell(8, 0).text = 'Длина отбортовки, мм'
    table.cell(8, 1).paragraphs[0].add_run('h_1').font.math = True
    #table.cell(8, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(8, 2).text =f'{data_in.dia}'
    table.add_row()
    table.cell(9, 0).text = 'Прибавка на коррозию, мм'
    table.cell(9, 1).paragraphs[0].add_run('c_1').font.math = True
    #table.cell(9, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(9, 2).text =f'{data_in.c_kor}'
    table.add_row()
    table.cell(10, 0).text = 'Прибавка для компенсации минусового допуска листа, мм'
    table.cell(10, 1).paragraphs[0].add_run('c_2').font.math = True
    #table.cell(10, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(10, 2).text =f'{data_in.c_minus}'
    table.add_row()
    table.cell(11, 0).text = 'Технологическая прибавка, мм'
    table.cell(11, 1).paragraphs[0].add_run('c_3').font.math = True
    #table.cell(11, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(11, 2).text =f'{data_in.c_3}'
    doc.add_paragraph('')
    doc.add_paragraph('Результаты расчета').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    doc.add_paragraph('Толщину стенки вычисляют по формуле:')
    doc.add_paragraph().add_run('s_1≥s_1p+c').font.math = True
    p = doc.add_paragraph('где ')
    p.add_run('s_1p').font.math = True
    p.add_run(' - расчетная толщина стенки днища')
    doc.add_paragraph().add_run('s_1p=(p∙R)/(2∙[σ]∙φ-0.5∙p)').font.math = True
    doc.add_paragraph('где R - радиус кривизны в вершине днища')
    # добавить расчет R для разных ситуаций
    if data_in.dia == data_out.elR:
        doc.add_paragraph(f'R=D={data_in.dia} мм - для эллиптичекских днищ с H=0.25D')
    else:
        doc.add_paragraph().add_run('R=D^2/(4∙H)').font.math = True
        doc.add_paragraph().add_run('R={data_in.dia}^2/(4∙{data_in.elH})={data_out.elR} мм').font.math = True
    doc.add_paragraph('c - сумма прибавок к расчетной толщине')
    doc.add_paragraph().add_run('c=c_1+c_2+c_3').font.math = True
    doc.add_paragraph().add_run(f'c={data_in.c_kor}+{data_in.c_minus}+{data_in.c_3}={data_out.c:.2f} мм').font.math = True
    doc.add_paragraph().add_run(f's_p=({data_in.press}∙{data_out.elR})/(2∙{data_in.sigma_d}∙{data_in.fi}-0.5{data_in.press})={data_out.s_calcr:.2f} мм').font.math = True
    doc.add_paragraph().add_run(f's={data_out.s_calcr:.2f}+{data_out.c:.2f}={data_out.s_calc:.2f} мм').font.math = True
    if data_in.s_prin >= data_out.s_calc:
        p = doc.add_paragraph(f'Принятая толщина ')
        p.add_run(f's_1={data_in.s_prin} мм').font.math = True
    else:
        doc.add_paragraph().add_run(f'Принятая толщина s={data_in.s_prin} мм').font.color.rgb = docx.shared.RGBColor(255,0,0)
    doc.add_paragraph('Допускаемое внутреннее избыточное давление вычисляют по формуле:')
    doc.add_paragraph().add_run('[p]=(2∙[σ]∙φ∙(s_1-c))/(R+0.5∙(s-c))').font.math = True
    doc.add_paragraph().add_run(f'[p]=(2∙{data_in.sigma_d}∙{data_in.fi}∙({data_in.s_prin}-{data_out.c:.2f}))/({data_out.elR}+0.5∙({data_in.s_prin}-{data_out.c:.2f}))={data_out.press_d:.2f} МПа').font.math = True
    doc.add_paragraph().add_run('[p]≥p').font.math = True
    doc.add_paragraph().add_run(f'{data_out.press_d:.2f}≥{data_in.press}').font.math = True
    if data_out.press_d > data_in.press:
        doc.add_paragraph('Условие прочности выполняется')
    else:
        doc.add_paragraph().add_run('Условие прочности не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)
    if data_out.ypf:
        p = doc.add_paragraph('Границы применения формул ')
    else:
        doc.add_paragraph().add_run('Границы применения формул не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)
    # эллептические днища
    doc.add_paragraph().add_run('0.002≤(s_1-c)/(D)≤0.1').font.math = True
    doc.add_paragraph().add_run(f'0.002≤({data_in.s_prin}-{data_out.c:.2f})/({data_in.dia})={(data_in.s_prin-data_out.c)/data_in.dia:.3f}≤0.1').font.math = True
    doc.add_paragraph().add_run('0.2≤H/D≤0.5').font.math = True
    doc.add_paragraph().add_run(f'0.2≤{data_in.elH}/{data_in.dia}={data_in.elH/data_in.dia:.3f}<0.5').font.math = True
    

    doc.save(docum)

def makeWord_elnar(data_in:CalcClass.data_in, data_out:CalcClass.data_out, docum=None):
    doc = docx.Document(docum)
    doc.add_heading(f'Расчет на прочность эллиптического днища {data_in.name}, нагруженного наружным давлением').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Исходные данные').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=0, cols=0)
    table.add_column(7000000)
    table.add_column(750000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.add_column(750000)
    table.add_row()  
    table.cell(0, 0).text = 'Наименование и размерность'
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = 'Обозначение'
    table.cell(0, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Значение'
    table.add_row()
    table.cell(1, 0).text = 'Расчетная температура, °С'
    table.cell(1, 1).text = 't'
    table.cell(1, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 2).text = f'{data_in.temp}'
    table.add_row()
    table.cell(2, 0).text = 'Расчетное внутреннее давление, МПа'
    table.cell(2, 1).text ='p'
    table.cell(2, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 2).text =f'{data_in.press}'
    table.add_row()
    table.cell(3, 0).text = 'Марка стали'
    table.cell(3, 1).text =''
    table.cell(3, 2).text =f'{data_in.steel}'
    table.add_row()
    table.cell(4, 0).text = 'Допускаемое напряжение при расчетной температуре, МПа'
    table.cell(4, 1).text = '[σ]'
    table.cell(4, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4, 2).text =f'{data_in.sigma_d}'
    table.add_row()
    table.cell(5, 0).text = 'Модуль продольной упругости при расчетной темп-ре, МПа'
    table.cell(5, 1).text = 'E'
    table.cell(5, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(5, 2).text =f'{data_in.E}'
    table.add_row()
    table.cell(6, 0).text = 'Коэффициент прочности сварного шва'
    table.cell(6, 1).paragraphs[0].add_run('φ_p').font.math = True
    #table.cell(6, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(6, 2).text =f'{data_in.fi}'
    table.add_row()
    table.cell(7, 0).text = 'Внутренний диаметр днища, мм'
    table.cell(7, 1).text = 'D'
    table.cell(7, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(7, 2).text =f'{data_in.dia}'
    table.add_row()
    table.cell(8, 0).text = 'Высота выпуклой части, мм'
    table.cell(8, 1).text = 'Н'
    table.cell(8, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(8, 2).text =f'{data_in.elH}'
    table.add_row()
    table.cell(9, 0).text = 'Длина отбортовки, мм'
    table.cell(9, 1).paragraphs[0].add_run('h_1').font.math = True
    #table.cell(9, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(9, 2).text =f'{data_in.dia}'
    table.add_row()
    table.cell(10, 0).text = 'Прибавка на коррозию, мм'
    table.cell(10, 1).paragraphs[0].add_run('c_1').font.math = True
    #table.cell(10, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(10, 2).text =f'{data_in.c_kor}'
    table.add_row()
    table.cell(11, 0).text = 'Прибавка для компенсации минусового допуска листа, мм'
    table.cell(11, 1).paragraphs[0].add_run('c_2').font.math = True
    #table.cell(11, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(11, 2).text =f'{data_in.c_minus}'
    table.add_row()
    table.cell(12, 0).text = 'Технологическая прибавка, мм'
    table.cell(12, 1).paragraphs[0].add_run('c_3').font.math = True
    #table.cell(12, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(12, 2).text =f'{data_in.c_3}'
    doc.add_paragraph('')
    doc.add_paragraph('Результаты расчета').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    doc.add_paragraph('Толщину стенки вычисляют по формуле:')
    doc.add_paragraph().add_run('s_1≥s_1p+c').font.math = True
    p = doc.add_paragraph('где ')
    p.add_run('s_1p').font.math = True
    p.add_run(' - расчетная толщина стенки днища')
    doc.add_paragraph().add_run('s_1p=max{(K_Э∙R)/(161)∙√((n_y∙p)/(10^-5∙E));(1.2∙p∙R)/(2∙[σ])}').font.math = True
    p = doc.add_paragraph('Для предварительного расчета ')
    p.add_run('К_Э=0.9').font.math = True
    p.add_run(' для эллиптических днищ')
    if data_in.dia == data_out.elR:
        doc.add_paragraph(f'R=D={data_in.dia} мм - для эллиптичекских днищ с H=0.25D')
    else:
        doc.add_paragraph().add_run('R=D^2/(4∙H)').font.math = True
        doc.add_paragraph().add_run('R={data_in.dia}^2/(4∙{data_in.elH})={data_out.elR} мм').font.math = True

    doc.add_paragraph().add_run(f'(0.9∙{data_out.elR})/(161)∙√(({data_in.ny}∙{data_in.press})/(10^-5∙{data_in.E}))={data_out.s_calcr1:.2f}').font.math = True
    doc.add_paragraph().add_run(f'(1.2∙{data_in.press}∙{data_out.elR})/(2∙{data_in.sigma_d})={data_out.s_calcr2:.2f}').font.math = True
    doc.add_paragraph().add_run(f's_1p=max({data_out.s_calcr1:.2f};{data_out.s_calcr2:.2f})={data_out.s_calcr:.2f} мм').font.math = True
        
    doc.add_paragraph('c - сумма прибавок к расчетной толщине')
    doc.add_paragraph().add_run('c=c_1+c_2+c_3').font.math = True
    doc.add_paragraph().add_run(f'c={data_in.c_kor}+{data_in.c_minus}+{data_in.c_3}={data_out.c:.2f} мм').font.math = True
        
    doc.add_paragraph().add_run(f's_1={data_out.s_calcr:.2f}+{data_out.c:.2f}={data_out.s_calc:.2f} мм').font.math = True
    if data_in.s_prin >= data_out.s_calc:
        p = doc.add_paragraph(f'Принятая толщина ')
        p.add_run(f's_1={data_in.s_prin} мм').font.math = True
    else:
        doc.add_paragraph().add_run(f'Принятая толщина s={data_in.s_prin} мм').font.color.rgb = docx.shared.RGBColor(255,0,0)

    doc.add_paragraph('Допускаемое наружное давление вычисляют по формуле:')
    doc.add_paragraph().add_run('[p]=[p]_П/√(1+([p]_П/[p]_E)^2)').font.math = True
    doc.add_paragraph('допускаемое давление из условия прочности вычисляют по формуле:')
    doc.add_paragraph().add_run('[p]_П=(2∙[σ]∙(s_1-c))/(R+0.5∙(s_1-c))').font.math = True
    doc.add_paragraph().add_run(f'[p]_П=(2∙{data_in.sigma_d}∙({data_in.s_prin}-{data_out.c:.2f}))/({data_out.elR}+0.5∙({data_in.s_prin}-{data_out.c:.2f}))={data_out.press_dp:.2f} МПа').font.math = True
    doc.add_paragraph('допускаемое давление из условия устойчивости в пределах упругости вычисляют по формуле:')
    doc.add_paragraph().add_run('[p]_E=(2.6∙10^-5∙E)/n_y∙[(100∙(s_1-c))/(K_Э∙R)]^2').font.math = True
    p = doc.add_paragraph('коэффициент ')
    p.add_run('K_Э').font.math = True
    p.add_run(' вычисляют по формуле')
    doc.add_paragraph().add_run('K_Э=(1+(2.4+8∙x)∙x)/(1+(3.0+10∙x)∙x)').font.math = True
    doc.add_paragraph().add_run('x=10∙(s_1-c)/D∙(D/(2∙H)-(2∙H)/D)').font.math = True
    doc.add_paragraph().add_run(f'x=10∙({data_in.s_prin}-{data_out.c:.2f})/{data_in.dia}∙({data_in.dia}/(2∙{data_in.elH})-(2∙{data_in.elH})/{data_in.dia})={data_out.elx:.2f}').font.math = True
    doc.add_paragraph().add_run(f'K_Э=(1+(2.4+8∙{data_out.elx:.2f})∙{data_out.elx:.2f})/(1+(3.0+10∙{data_out.elx:.2f})∙{data_out.elx:.2f})={data_out.elke:.2f}').font.math = True
    
    doc.add_paragraph().add_run(f'[p]_E=(2.6∙10^-5∙{data_in.E})/{data_in.ny}∙[(100∙({data_in.s_prin}-{data_out.c:.2f}))/({data_out.elke:.2f}∙{data_out.elR})]^2={data_out.press_de:.2f} МПа').font.math = True
    doc.add_paragraph().add_run(f'[p]={data_out.press_dp:.2f}/√(1+({data_out.press_dp:.2f}/{data_out.press_de:.2f})^2)={data_out.press_d:.2f} МПа').font.math = True
    doc.add_paragraph().add_run('[p]≥p').font.math = True
    doc.add_paragraph().add_run(f'{data_out.press_d:.2f}≥{data_in.press}').font.math = True
    if data_out.press_d > data_in.press:
        doc.add_paragraph('Условие прочности выполняется')
    else:
        doc.add_paragraph().add_run('Условие прочности не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)
    
    if data_out.ypf:
        p = doc.add_paragraph('Границы применения формул ')
    else:
        doc.add_paragraph().add_run('Границы применения формул не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)
    # эллептические днища
    doc.add_paragraph().add_run('0.002≤(s_1-c)/(D)≤0.1').font.math = True
    doc.add_paragraph().add_run(f'0.002≤({data_in.s_prin}-{data_out.c:.2f})/({data_in.dia})={(data_in.s_prin-data_out.c)/data_in.dia:.3f}≤0.1').font.math = True
    doc.add_paragraph().add_run('0.2≤H/D≤0.5').font.math = True
    doc.add_paragraph().add_run(f'0.2≤{data_in.elH}/{data_in.dia}={data_in.elH/data_in.dia:.3f}<0.5').font.math = True

    doc.save(docum)

def makeWord_obvnyk(data_in:CalcClass.data_in, data_out:CalcClass.data_out, data_nozzlein:CalcClass.data_nozzlein, data_nozzleout:CalcClass.data_nozzleout, docum=None):
    doc = docx.Document(docum)
    doc.add_heading(f'Расчет на прочность узла врезки штуцера {data_nozzlein.name} в обечайку {data_in.name}, нагруженную внутренним избыточным давлением').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Исходные данные').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=0, cols=0)
    table.add_column(3000000)
    table.add_column(3000000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.add_row()  
    table.cell(0, 0).text = 'Элемент:'
    #table.cell(0, 0).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = f'Штуцер {data_nozzlein.name}'
    #table.cell(0, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.add_row()
    table.cell(1, 0).text = 'Элемент несущий штуцер:'
    table.cell(1, 1).text = f'{data_in.name}'
    table.add_row()
    table.cell(2, 0).text = 'Тип элемента, несущего штуцер'
    if data_in.met == 'obvn' or data_in.met == 'obnar':
        table.cell(2, 1).text ='Обечайка цилиндрическая'
    elif data_in.met == 'konvn' or data_in.met == 'konnar':
        table.cell(2, 1).text ='Обечайка коническая'
    elif data_in.met == 'elvn' or data_in.met == 'elnar':
        table.cell(2, 1).text ='Днище эллиптическое'
    table.add_row()
    table.cell(3, 0).text = 'Тип штуцера'
    if data_nozzlein.vid == 1:
        table.cell(3, 1).text ='Непроходящий без укрепления'
    elif data_nozzlein.vid == 2:
        table.cell(3, 1).text ='Проходящий без укрепления'
    elif data_nozzlein.vid == 3:
        table.cell(3, 1).text ='Непроходящий с накладным кольцом'
    elif data_nozzlein.vid == 4:
        table.cell(3, 1).text ='Проходящий с накладным кольцом'
    elif data_nozzlein.vid == 5:
        table.cell(3, 1).text ='С накладным кольцом и внутренней частью'
    elif data_nozzlein.vid == 6:
        table.cell(3, 1).text ='С отбортовкой'
    elif data_nozzlein.vid == 7:
        table.cell(3, 1).text ='С торовой вставкой'
    elif data_nozzlein.vid == 8:
        table.cell(3, 1).text ='С вварным кольцом'
    doc.add_picture(f'pic/Nozzle/Nozzle{data_nozzlein.vid}.gif', width=docx.shared.Inches(3))
    
      
    table = doc.add_table(rows=0, cols=0)
    table.add_column(5000000)
    table.add_column(3000000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.add_row()  
    table.cell(0, 0).text = 'Материал несущего элемента:'
    table.cell(0, 1).text = f'{data_in.steel}'
    table.add_row()
    table.cell(1, 0).text = 'Толщина стенки несущего элемента, s:'
    table.cell(1, 1).text = f'{data_in.s_prin} мм'
    table.add_row()
    table.cell(2, 0).text = 'Сумма прибавок к стенке несущего элемента, c:'
    table.cell(2, 1).text =f'{data_out.c:.2f} мм'
    table.add_row()
    table.cell(3, 0).text = 'Материал штуцера'
    table.cell(3, 1).text =f'{data_nozzlein.steel1}'
    table.add_row()
    table.cell(4, 0).text = 'Внутренний диаметр штуцера, d'
    table.cell(4, 1).text =f'{data_nozzlein.dia} мм'
    table.add_row()
    table.cell(5, 0).text = 'Толщина стенки штуцера, '
    table.cell(5, 0).paragraphs[0].add_run('s_1').font.math = True
    table.cell(5, 1).text =f'{data_nozzlein.s1} мм'
    table.add_row()
    table.cell(6, 0).text = 'Длина наружной части штуцера, '
    table.cell(6, 0).paragraphs[0].add_run('l_1').font.math = True
    table.cell(6, 1).text =f'{data_nozzlein.l1} мм'
    table.add_row()
    table.cell(7, 0).text = 'Сумма прибавок к толщине стенки штуцера, '
    table.cell(7, 0).paragraphs[0].add_run('c_s').font.math = True
    table.cell(7, 1).text =f'{data_nozzlein.cs}'
    table.add_row()
    table.cell(8, 0).text = 'Прибавка на коррозию, '
    table.cell(8, 0).paragraphs[0].add_run('c_s1').font.math = True
    table.cell(8, 1).text =f'{data_nozzlein.cs1}'
    if data_nozzlein.vid == 1:
        table.add_row()
        table.cell(9, 0).text = 'Минимальный размер сварного шва, Δ'
        table.cell(9, 1).text =f'{data_nozzlein.delta} мм'
    elif data_nozzlein.vid == 2:
        table.add_row()
        table.cell(9, 0).text = 'Длина внутренней части штуцера, '
        table.cell(9, 0).paragraphs[0].add_run('l_3').font.math = True
        table.cell(9, 1).text =f'{data_nozzlein.l1} мм'
        table.add_row()
        table.cell(10, 0).text = 'Минимальный размер сварного шва, Δ'
        table.cell(10, 1).text =f'{data_nozzlein.delta} мм'
    elif data_nozzlein.vid == 3:
        table.add_row()
        table.cell(9, 0).text = 'Ширина накладного кольца, '
        table.cell(9, 0).paragraphs[0].add_run('l_2').font.math = True
        table.cell(9, 1).text =f'{data_nozzlein.l2} мм'
        table.add_row()
        table.cell(10, 0).text = 'Толщина накладного кольца, '
        table.cell(10, 0).paragraphs[0].add_run('s_2').font.math = True
        table.cell(10, 1).text =f'{data_nozzlein.s2} мм'
        table.add_row()
        table.cell(11, 0).text = 'Минимальный размер сварного шва, Δ'
        table.cell(11, 1).text =f'{data_nozzlein.delta} мм'
    elif data_nozzlein.vid == 4:
        table.add_row()
        table.cell(9, 0).text = 'Ширина накладного кольца, '
        table.cell(9, 0).paragraphs[0].add_run('l_2').font.math = True
        table.cell(9, 1).text =f'{data_nozzlein.l2} мм'
        table.add_row()
        table.cell(10, 0).text = 'Толщина накладного кольца, '
        table.cell(10, 0).paragraphs[0].add_run('s_2').font.math = True
        table.cell(10, 1).text =f'{data_nozzlein.s2} мм'
        table.add_row()
        table.cell(11, 0).text = 'Длина внутренней части штуцера, '
        table.cell(11, 0).paragraphs[0].add_run('l_3').font.math = True
        table.cell(11, 1).text =f'{data_nozzlein.l1} мм'
        table.add_row()
        table.cell(12, 0).text = 'Минимальный размер сварного шва, Δ'
        table.cell(12, 1).text =f'{data_nozzlein.delta} мм'
    elif data_nozzlein.vid == 5:
        table.add_row()
        table.cell(9, 0).text = 'Ширина накладного кольца, '
        table.cell(9, 0).paragraphs[0].add_run('l_2').font.math = True
        table.cell(9, 1).text =f'{data_nozzlein.l2} мм'
        table.add_row()
        table.cell(10, 0).text = 'Толщина накладного кольца, '
        table.cell(10, 0).paragraphs[0].add_run('s_2').font.math = True
        table.cell(10, 1).text =f'{data_nozzlein.s2} мм'
        table.add_row()
        table.cell(11, 0).text = 'Длина внутренней части штуцера, '
        table.cell(11, 0).paragraphs[0].add_run('l_3').font.math = True
        table.cell(11, 1).text =f'{data_nozzlein.l1} мм'
        table.add_row()
        table.cell(12, 0).text = 'Толщина внутренней части штуцера, '
        table.cell(12, 0).paragraphs[0].add_run('s_3').font.math = True
        table.cell(12, 1).text =f'{data_nozzlein.s3} мм'
        table.add_row()
        table.cell(13, 0).text = 'Минимальный размер сварного шва, Δ'
        table.cell(13, 1).text =f'{data_nozzlein.delta} мм'
    elif data_nozzlein.vid == 6:
        table.add_row()
        table.cell(9, 0).text = 'Радиус отбортовки, r'
        #table.cell(9, 1).text =f'{data_nozzlein.rotbort} мм'
        table.add_row()
        table.cell(10, 0).text = 'Минимальный размер сварного шва, Δ'
        table.cell(10, 1).text =f'{data_nozzlein.delta} мм'
    elif data_nozzlein.vid == 7:
        pass
    elif data_nozzlein.vid == 8:
        pass

    doc.add_paragraph('Коэффициенты прочности сварных швов:')
    p = doc.add_paragraph('Продольный шов штуцера ')
    p.add_run(f'φ_1={data_nozzlein.fi1}').font.math = True
    p = doc.add_paragraph('Шов обечайки в зоне врезки штуцера ')
    p.add_run(f'φ={data_nozzlein.fi}').font.math = True

    doc.add_paragraph('')
    doc.add_paragraph('Условия нагружения').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    table = doc.add_table(rows=0, cols=0)
    table.add_column(5000000)
    table.add_column(3000000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.add_row()  
    table.cell(0, 0).text = 'Расчетная температура, Т:'
    table.cell(0, 1).text = f'{data_in.temp} °С'
    table.add_row()
    if data_in.dav == 'vn':
        table.cell(1, 0).text = 'Расчетное внутреннее избыточное давление, p:'
    else:
        table.cell(1, 0).text = 'Расчетное наружное давление, p:'
    table.cell(1, 1).text = f'{data_in.press} МПа'

    table.add_row()
    table.cell(2, 0).text = f'Допускаемое напряжение для материала {data_nozzlein.steel1} при расчетной температуре, '
    table.cell(2, 0).paragraphs[0].add_run('[σ]_1').font.math = True
    table.cell(2, 0).paragraphs[0].add_run(':').font.math = True
    table.cell(2, 1).text =f'{data_nozzlein.sigma_d1} МПа'
    if data_in.dav == 'nar':
        table.add_row()
        table.cell(3, 0).text = 'Модуль продольной упругости при расчетной температуре, '
        table.cell(3, 0).paragraphs[0].add_run('E_1').font.math = True
        table.cell(3, 0).paragraphs[0].add_run(':').font.math = True
        table.cell(3, 1).text =f'{data_nozzlein.E1} МПа'
    if data_nozzlein.steel1 != data_nozzlein.steel2:
        table = doc.add_table(rows=0, cols=0)
        table.add_column(5000000)
        table.add_column(3000000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.add_row()
        table.cell(0, 0).text = 'Допускаемое напряжение для материала {data_nozzlein.steel2} при расчетной температуре, '
        table.cell(0, 0).paragraphs[0].add_run('[σ]_2').font.math = True
        table.cell(0, 0).paragraphs[0].add_run(':').font.math = True
        table.cell(0, 1).text =f'{data_nozzlein.sigma_d2} МПа'
        if data_in.dav == 'nar':
            table.add_row()
            table.cell(1, 0).text = 'Модуль продольной упругости при расчетной температуре, '
            table.cell(1, 0).paragraphs[0].add_run('E_2').font.math = True
            table.cell(1, 0).paragraphs[0].add_run(':').font.math = True
            table.cell(1, 1).text =f'{data_nozzlein.E2} МПа'
    if data_nozzlein.steel1 != data_nozzlein.steel3:
        table = doc.add_table(rows=0, cols=0)
        table.add_column(5000000)
        table.add_column(3000000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.add_row()
        table.cell(0, 0).text = 'Допускаемое напряжение для материала {data_nozzlein.steel3} при расчетной температуре, '
        table.cell(0, 0).paragraphs[0].add_run('[σ]_3').font.math = True
        table.cell(0, 0).paragraphs[0].add_run(':').font.math = True
        table.cell(0, 1).text =f'{data_nozzlein.sigma_d3} МПа'
        if data_in.dav == 'nar':
            table.add_row()
            table.cell(1, 0).text = 'Модуль продольной упругости при расчетной температуре, '
            table.cell(1, 0).paragraphs[0].add_run('E_3').font.math = True
            table.cell(1, 0).paragraphs[0].add_run(':').font.math = True
            table.cell(1, 1).text =f'{data_nozzlein.E3} МПа'
    
    doc.add_paragraph('')
    doc.add_paragraph('Результаты расчета').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
        
    doc.add_paragraph('Расчетные параметры').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    doc.add_paragraph('Расчетный диаметр укрепляемого элемента (для цилиндрической обечайки)')
    doc.add_paragraph().add_run(f'D_p=D={data_in.dia} мм').font.math = True 
    
    doc.add_paragraph('Расчетный диаметр отверстия в стенке цилиндрической обечайки при наличии штуцера с круглым поперечным сечением, ось которого совпадает с нормалью к поверхности в центре отверстия')
    doc.add_paragraph().add_run('d_p=d+2∙c_s').font.math = True
    doc.add_paragraph().add_run(f'd_p={data_nozzlein.dia}+2∙{data_nozzlein.cs}={data_nozzleout.dp:.2f} мм').font.math = True

    doc.add_paragraph('Расчетная длина внешней части штуцера')
    doc.add_paragraph().add_run('l_1p=min{l_1;1.25√((d+2∙c_s)(s_1-c_s))}').font.math = True
    doc.add_paragraph().add_run(f'1.25√((d+2∙c_s)(s_1-c_s))=1.25√(({data_nozzlein.dia}+2∙{data_nozzlein.cs})({data_nozzlein.s1}-{data_nozzlein.cs}))={data_nozzleout.l1p2:.2f} мм').font.math = True
    doc.add_paragraph().add_run(f'l_1p=min({data_nozzlein.l1};{data_nozzleout.l1p2:.2f})={data_nozzleout.l1p:.2f} мм').font.math = True

    if data_nozzlein.l3 > 0:
        doc.add_paragraph('Расчетная длина внешней части штуцера')
        doc.add_paragraph().add_run('l_3p=min{l_3;0.5√((d+2∙c_s)(s_3-c_s-c_s1))}').font.math = True
        doc.add_paragraph().add_run(f'0.5√((d+2∙c_s)(s_3-c_s-c_s1))=0.5√(({data_nozzlein.dia}+2∙{data_nozzlein.cs})({data_nozzlein.s3}-{data_nozzlein.cs}-{data_nozzlein.cs1}))={data_nozzleout.l3p2:.2f} мм').font.math = True
        doc.add_paragraph().add_run(f'l_3p=min({data_nozzlein.l3};{data_nozzleout.l3p2:.2f})={data_nozzleout.l3p:.2f} мм').font.math = True
    
    doc.add_paragraph('Ширина зоны укрепления отверстия в цилиндрической обечайке')
    doc.add_paragraph().add_run('L_0=√(D_p∙(s-c))').font.math = True
    doc.add_paragraph().add_run(f'L_0=√({data_nozzleout.Dp}∙({data_in.s_prin}-{data_out.c:.2f}))={data_nozzleout.L0:.2f}').font.math = True

    doc.add_paragraph('Расчетная ширина зоны укрепления отверстия в стенке циллиндрической обечайки')
    if data_nozzlein.vid in [1, 2, 3, 4, 5, 6]:
        doc.add_paragraph().add_run(f'l_p=L_0={data_nozzleout.lp:.2f} мм').font.math = True
    elif data_nozzlein.vid in [7, 8]:
        doc.add_paragraph().add_run('l_p=min{l;L_0}').font.math = True
        doc.add_paragraph().add_run(f'l_p=min({data_nozzlein.l};{data_nozzleout.L0:.2f})={data_nozzleout.lp:.2f} мм').font.math = True

    if data_nozzlein.l2 > 0:
        doc.add_paragraph('Расчетная ширина накладного кольца')
        doc.add_paragraph().add_run('l_2p=min{l_2;√(D_p∙(s_2+s-c))}').font.math = True
        doc.add_paragraph().add_run(f'√(D_p∙(s_2+s-c))=√({data_nozzleout.Dp}∙({data_nozzlein.s2}+{data_in.s_prin}-{data_out.c:.2f}))={data_nozzleout.l2p2:.2f} мм').font.math = True
        doc.add_paragraph().add_run(f'l_2p=min({data_nozzlein.l2};{data_nozzleout.l2p2:.2f})={data_nozzleout.l2p:.2f} мм').font.math = True

    doc.add_paragraph('Учет применения различного материального исполнения')
    if data_in.steel != data_nozzlein.steel1:
        p = doc.add_paragraph('- для внешней части штуцера')
        p.add_run(f'χ_1=min(1;[σ]_1/[σ])=min(1;{data_nozzlein.sigmad1}/{data_in.sigma_d})={data_nozzleout.psi1}')
    if data_in.steel != data_nozzlein.steel2:
        p = doc.add_paragraph('- для накладного кольца')
        p.add_run(f'χ_2=min(1;[σ]_2/[σ])=min(1;{data_nozzlein.sigmad2}/{data_in.sigma_d})={data_nozzleout.psi2}')
    if data_in.steel != data_nozzlein.steel3:
        p = doc.add_paragraph('- для внутренней части штуцера')
        p.add_run(f'χ_3=min(1;[σ]_3/[σ])=min(1;{data_nozzlein.sigmad3}/{data_in.sigma_d})={data_nozzleout.psi3}')
    if data_in.steel != data_nozzlein.steel4:
        p = doc.add_paragraph('- для торообразной вставки или вварного кольца')
        p.add_run(f'χ_4=min(1;[σ]_4/[σ])=min(1;{data_nozzlein.sigmad4}/{data_in.sigma_d})={data_nozzleout.psi4}')

    doc.add_paragraph('Расчетный диаметр отверстия, не требующий укрепления в стенке цилиндрической обечайки при отсутствии избыточной толщины стенки сосуда и при наличии штуцера')
    doc.add_paragraph().add_run('d_0p=0,4√(D_p∙(s-c))').font.math = True
    doc.add_paragraph().add_run(f'd_0p=0.4√({data_nozzleout.Dp}∙({data_in.s_prin}-{data_out.c:.2f}))={data_nozzleout.d0p:.2f} мм').font.math = True

            
    doc.add_paragraph('Проверка условия необходимости проведения расчета укрепления отверстий')
    doc.add_paragraph().add_run('d_p≤d_0').font.math = True
    
    p = doc.add_paragraph('')
    p.add_run('d_0').font.math = True
    p.add_run(' - наибольший допустимый диаметр одиночного отверстия, не требующего дополнительного укрепления при наличии избыточной толщины стенки сосуда')
    doc.add_paragraph().add_run('d_0=min{2∙((s-c)/s_pn-0.8)∙√(D_p∙(s-c));d_max+2∙c_s}').font.math = True
    
    p = doc.add_paragraph('где ')
    p.add_run('d_max').font.math = True
    p.add_run(' - максимальный диаметр отверстия ')

    p = doc.add_paragraph('')
    p.add_run(f'd_max=D={data_in.dia} мм').font.math = True 
    p.add_run(' - для отверстий в цилиндрических обечайках')
    
    
    
    if data_in.dav == 'vn':
        p = doc.add_paragraph('')
        p.add_run(f's_pn=s_p={data_out.s_calcr:.2f} мм').font.math = True 
        p.add_run(' - в случае внутреннего давления')
    else:
        doc.add_paragraph().add_run('s_pn=(p_pn∙D_p)/(2∙K_1∙[σ]-p_pn)').font.math = True
        doc.add_paragraph().add_run('p_pn=p/√(1-(p/[p]_E)^2)').font.math = True
        p = doc.add_paragraph('')
        p.add_run('[p]_E').font.math = True
        p.add_run(' -  допускаемое наружное давление из условия устойчивости в пределах упругости, определяемое по ГОСТ 34233.2 для обечайки без отверстий')
        doc.add_paragraph().add_run(f'p_pn={data_in.press}/√(1-({data_in.press}/{data_nozzleout.pen:.2f})^2)={data_nozzleout.ppn:.2f} МПа').font.math = True
        doc.add_paragraph().add_run('s_pn=({data_nozzleout.ppn:.2f}∙{data_nozzleout.Dp:.2f})/(2∙{data_nozzleout.K1}∙{data_in.sigma_d}-{data_nozzleout.ppn:.2f})={data_nozzleout.spn:.2f} мм').font.math = True

    doc.add_paragraph().add_run(f'2∙((s-c)/s_pn-0.8)∙√(D_p∙(s-c))=2∙(({data_in.s_prin}-{data_out.c:.2f})/{data_nozzleout.spn:.2f}-0.8)∙√({data_nozzleout.Dp}∙({data_in.s_prin}-{data_out.c:.2f}))={data_nozzleout.d01:.2f}').font.math = True
    doc.add_paragraph().add_run(f'd_max+2∙c_s={data_nozzleout.dmax:.2f}+2∙{data_nozzlein.cs}={data_nozzleout.d02:.2f}').font.math = True
    doc.add_paragraph().add_run(f'd_0=min({data_nozzleout.d01:.2f};{data_nozzleout.d02:.2f})={data_nozzleout.d0:.2f} мм').font.math = True
    if data_nozzleout.dp <= data_nozzleout.d0:
        doc.add_paragraph().add_run(f'{data_nozzleout.dp:.2f}≤{data_nozzleout.d0:.2f}').font.math = True
        doc.add_paragraph('Условие прочности выполняется, следовательно дальнейших расчетов укрепления отверстия не требуется')
    else:
        doc.add_paragraph().add_run(f'{data_nozzleout.dp:.2f}≤{data_nozzleout.d0:.2f}').font.math = True
        doc.add_paragraph('Условие прочности не выполняется, следовательно необходим дальнейший расчет укрепления отверстия')
        doc.add_paragraph('В случае укрепления отверстия утолщением стенки сосуда или штуцера, или накладным кольцом, или вварным кольцом, или торообразной вставкой, или отбортовкой должно выполняться условие')
        doc.add_paragraph().add_run('l_1p∙(s_1-s_1p-c_s)∙χ_1+l_2p∙s_2∙χ_2+l_3p∙(s_3-c_s-c_s1)∙χ_3+l1p∙(s-s_p-c)∙χ_4≥0.5∙(d_p-d_0p)∙s_p').font.math = True
        doc.add_paragraph().add_run('l_1p∙(s_1-s_1p-c_s)∙χ_1+l_2p∙s_2∙χ_2+l_3p∙(s_3-c_s-c_s1)∙χ_3+l1p∙(s-s_p-c)∙χ_4=').font.math = True
        doc.add_paragraph().add_run(f'{data_nozzleout.l1p:.2f}∙({data_nozzlein.s1}-{data_nozzleout.s1p:.2f}-{data_nozzlein.cs})∙{data_nozzleout.psi1}+{data_nozzleout.l2p:.2f}∙{data_nozzlein.s2}∙{data_nozzleout.psi2}+{data_nozzleout.l3p:.2f}∙({data_nozzlein.s3}-{data_nozzlein.cs}-{data_nozzlein.cs1})∙{data_nozzleout.psi3:.2f}+{data_nozzleout.lp:.2f}∙({data_in.s_prin}-{data_out.s_calcr:.2f}-{data_out.c:.2f})∙{data_nozzleout.psi4}={data_nozzleout.yslyk1:.2f}').font.math = True
        doc.add_paragraph().add_run(f'0.5∙(d_p-d_0p)∙s_p=0.5∙({data_nozzleout.dp:.2f}-{data_nozzleout.d0p:.2f})∙{data_out.s_calcr:.2f}={data_nozzleout.yslyk2:.2f}').font.math = True
        doc.add_paragraph().add_run(f'{data_nozzleout.yslyk1:.2f}≥{data_nozzleout.yslyk2:.2f}').font.math = True
        if data_nozzleout.yslyk1 >= data_nozzleout.yslyk2:
            doc.add_paragraph('Условие прочности выполняется')
        else:
            doc.add_paragraph().add_run('Условие прочности не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)


    p = doc.add_paragraph('')
    doc.add_paragraph('Допускаемое внутреннее избыточное давление элемента сосуда с учетом ослабления стенки отверстием вычисляют по формуле')
    doc.add_paragraph().add_run('[p]=(2∙K_1∙φ∙[σ]∙(s-c)∙V)/(D_p+(s-c)∙V)').font.math = True

    p = doc.add_paragraph('где ')
    p.add_run('K_1=1').font.math = True
    p.add_run(' - для цилиндрических и конических обечаек')
    doc.add_paragraph('Коэффициент снижения прочности сосуда, ослабленного одиночным отверстием, вычисляют по формуле')
    doc.add_paragraph().add_run('V=min{(s_0-c)/(s-c);(χ_4+(l_1p∙(s_1-c_s)∙χ_1+l_2p∙s_2∙χ_2+l_3p∙(s_3-c_s-c_s1)∙χ_3)/(l_p∙(s-c)))/(1+0.5∙(d_p-d_0p)/l_p+K_1∙(d+2∙c_s)/D_p∙(φ/φ_1)∙(l_1p/l_p))}').font.math = True
    #doc.add_picture(f'pic/Nozzle/V.png', width=docx.shared.Inches(6))
    l = doc.paragraphs[-1]
    l.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    if data_nozzlein.vid in [1, 2, 6]:
        p = doc.add_paragraph('При отсутствии накладного кольца и укреплении отверстия штуцером ')
        p.add_run('s_2=0 , s_0=s , χ_4=1').font.math = True
    elif data_nozzlein.vid in [3, 4, 5]:
        p = doc.add_paragraph('При отсутствии вварного кольца или торообразной вставки ')
        p.add_run('s_0=s , χ_4=1').font.math = True
    


    doc.add_paragraph().add_run(f'(s_0-c)/(s-c)=({data_nozzlein.s0}-{data_out.c:.2f})/({data_in.s_prin}-{data_out.c:.2f})={data_nozzleout.V1:.2f}').font.math = True
    doc.add_paragraph().add_run('(χ_4+(l_1p∙(s_1-c_s)∙χ_1+l_2p∙s_2∙χ_2+l_3p∙(s_3-c_s-c_s1)∙χ_3)/(l_p∙(s-c)))/(1+0.5∙(d_p-d_0p)/l_p+K_1∙(d+2∙c_s)/D_p∙(φ/φ_1)∙(l_1p/l_p))=').font.math = True
    doc.add_paragraph().add_run(f'({data_nozzleout.psi4}+({data_nozzleout.l1p:.2f}∙({data_nozzlein.s1}-{data_nozzlein.cs})∙{data_nozzleout.psi1}+{data_nozzleout.l2p:.2f}∙{data_nozzlein.s2}∙{data_nozzleout.psi2}+{data_nozzleout.l3p:.2f}∙({data_nozzlein.s3}-{data_nozzlein.cs}-{data_nozzlein.cs1})∙{data_nozzleout.psi3:.2f})/({data_nozzleout.lp:.2f}∙({data_in.s_prin}-{data_out.c:.2f})))/(1+0.5∙({data_nozzleout.dp:.2f}-{data_nozzleout.d0p:.2f})/{data_nozzleout.lp:.2f}+{data_nozzleout.K1}∙({data_nozzlein.dia}+2∙{data_nozzlein.cs})/{data_nozzleout.Dp}∙({data_nozzlein.fi}/{data_nozzlein.fi1})∙({data_nozzleout.l1p:.2f}/{data_nozzleout.lp:.2f}))={data_nozzleout.V2:.2f}').font.math = True

    doc.add_paragraph().add_run(f'V=min({data_nozzleout.V1:.2f};{data_nozzleout.V2:.2f})={data_nozzleout.V:.2f} мм').font.math = True

    doc.add_paragraph().add_run(f'[p]=(2∙{data_nozzleout.K1}∙{data_nozzlein.fi}∙{data_in.sigma_d}∙({data_in.s_prin}-{data_out.c:.2f})∙{data_nozzleout.V:.2f})/({data_nozzleout.Dp}+({data_in.s_prin}-{data_out.c:.2f})∙{data_nozzleout.V:.2f})={data_nozzleout.press_d:.2f} МПа').font.math = True
    doc.add_paragraph().add_run('[p]≥p')
    doc.add_paragraph().add_run(f'{data_nozzleout.press_d:.2f} МПа >= {data_in.press} МПа').font.math = True
    if data_nozzleout.press_d >= data_in.press:
        doc.add_paragraph('Условие прочности выполняется')
    else:
        doc.add_paragraph().add_run('Условие прочности не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)
    
    doc.add_paragraph('Границы применения формул')
    doc.add_paragraph().add_run('(d_p-2∙c_s)/D≤1').font.math = True
    doc.add_paragraph().add_run(f'({data_nozzleout.dp:.2f}-2∙{data_nozzlein.cs})/{data_in.dia}={(data_nozzleout.dp-2*data_nozzlein.cs)/data_in.dia:.2f}≤1').font.math = True
    doc.add_paragraph().add_run('(s-c)/D≤0.1').font.math = True
    doc.add_paragraph().add_run(f'({data_in.s_prin}-{data_out.c:.2f})/({data_in.dia})={(data_in.s_prin-data_out.c)/data_in.dia:.3f}≤0.1').font.math = True

    doc.save(docum)

def makeWord_obsaddle(data_in:CalcClass.data_saddlein, data_out:CalcClass.data_saddleout, docum=None):
    doc = docx.Document(docum)

    doc = docx.Document(docum)
    doc.add_heading(f'Расчет на прочность обечайки {data_in.nameob} от воздействия опорных нагрузок').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Исходные данные').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    if data_in.type == 1:
        doc.add_picture('pic/Saddle/SaddleNothingElem.gif')
    elif data_in.type == 2:
        doc.add_picture('pic/Saddle/SaddleSheetElem.gif')
    
    table = doc.add_table(rows=0, cols=0)
    table.add_column(7000000)
    table.add_column(750000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.add_column(750000)
    table.add_row() 
    table.cell(0, 0).text = 'Наименование и размерность'
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).text = 'Обозначение'
    table.cell(0, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Значение'
    table.add_row()  
    table.cell(1, 0).text = 'Внутренний диаметр обечайки, мм'
    table.cell(1, 1).text = 'D'
    table.cell(1, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 2).text =f'{data_in.D}'
    table.add_row()
    table.cell(2, 0).text = 'Толщина стенки обечайки, мм'
    table.cell(2, 1).text = 's'
    table.cell(2, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(2, 2).text = f'{data_in.s}'
    table.add_row()
    table.cell(3, 0).text = 'Прибавка к расчетной толщине, мм'
    table.cell(3, 1).text = 'c'
    table.cell(3, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(3, 2).text =f'{data_in.c}'
    table.add_row()
    table.cell(4, 0).text = 'Длина обечайки, мм'
    table.cell(4, 1).text = 'Lob'
    table.cell(4, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4, 2).text =f'{data_in.Lob}'
    table.add_row()
    table.cell(5, 0).text = 'Коэффициент прочности сварного шва'
    table.cell(5, 1).text = 'φ'
    table.cell(5, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(5, 2).text =f'{data_in.fi}'
    table.add_row()
    table.cell(6, 0).text = 'Марка стали'
    table.cell(6, 1).text =''
    table.cell(6, 2).text =f'{data_in.steel}'
    table.add_row()
    table.cell(7, 0).text = 'Собственный вес с содержимым, Н'
    table.cell(7, 1).text ='G'
    table.cell(7, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(7, 2).text =f'{data_in.G}'
    table.add_row()
    table.cell(8, 0).text = 'Ширина опоры, мм'
    table.cell(8, 1).text = 'b'
    table.cell(8, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(8, 2).text =f'{data_in.b}'
    table.add_row()
    table.cell(9, 0).text = 'Угол охвата опоры, '
    table.cell(9, 1).paragraphs[0].add_run('δ_1').font.math = True
    #table.cell(9, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(9, 2).text =f'{data_in.delta1}'
    table.add_row()
    table.cell(10, 0).text = 'Длина свободно выступающей части, мм'
    table.cell(10, 1).text = 'e'
    table.cell(10, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(10, 2).text =f'{data_in.e:.2f}'
    table.add_row()
    table.cell(11, 0).text = 'Длина выступающей цилиндрической части сосуда, включая отбортовку днища, мм'
    table.cell(11, 1).text = 'a'
    table.cell(11, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(11, 2).text =f'{data_in.a}'
    table.add_row()
    table.cell(12, 0).text = 'Высота опоры, мм'
    table.cell(12, 1).text = 'H'
    table.cell(12, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table.cell(12, 2).text =f'{data_in.H}'
    if data_in.type == 2:
        table.add_row()
        table.cell(13, 0).text = 'Толщина подкладного листа, мм'
        table.cell(13, 1).paragraphs[0].add_run('s_2').font.math = True
        table.cell(13, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        table.cell(13, 2).text = f'{data_in.s2}'
        table.add_row()
        table.cell(14, 0).text = 'Ширина подкладного листа, мм'
        table.cell(14, 1).paragraphs[0].add_run('b_2').font.math = True
        #table.cell(14, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        table.cell(14, 2).text =f'{data_in.b2}'
        table.add_row()
        table.cell(15, 0).text = 'Угол охвата подкладного листа, '
        table.cell(15, 1).paragraphs[0].add_run('δ_2').font.math = True
        #table.cell(15, 1).paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        table.cell(15, 2).text =f'{data_in.delta2}'
    doc.add_paragraph('')

    doc.add_paragraph('Условия нагружения').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    table = doc.add_table(rows=0, cols=0)
    table.add_column(5000000)
    table.add_column(3000000) #.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.add_row()  
    table.cell(0, 0).text = 'Расчетная температура, Т:'
    table.cell(0, 1).text = f'{data_in.temp} °С'
    table.add_row()
    if data_in.dav == 'vn':
        table.cell(1, 0).text = 'Расчетное внутреннее избыточное давление, p:'
    else:
        table.cell(1, 0).text = 'Расчетное наружное давление, p:'
    table.cell(1, 1).text = f'{data_in.p} МПа'
    table.add_row()
    table.cell(2, 0).text = f'Допускаемое напряжение для материала {data_in.steel} при расчетной температуре, [σ]:'
    table.cell(2, 1).text =f'{data_in.sigma_d} МПа'
    table.add_row()
    table.cell(3, 0).text = 'Модуль продольной упругости при расчетной температуре, E:'
    table.cell(3, 1).text =f'{data_in.E} МПа'

    doc.add_paragraph('')
    doc.add_paragraph('Результаты расчета').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
        
    doc.add_paragraph('Расчетные параметры').paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    doc.add_paragraph('Распределенная весовая нагрузка')
    doc.add_paragraph().add_run('q=G/(L+4/3∙H)').font.math = True 
    doc.add_paragraph().add_run(f'q={data_in.G}/({data_in.L}+4/3∙{data_in.H})={data_out.q:.2f} Н/мм').font.math = True 

    doc.add_paragraph('Расчетный изгибающий момент, действующий на консольную часть обечайки')
    doc.add_paragraph().add_run('M_0=q∙D^2/16').font.math = True 
    doc.add_paragraph().add_run(f'M_0={data_out.q:.2f}∙{data_in.D}^2/16={data_out.M0:.2f} Н∙мм').font.math = True 

    doc.add_paragraph('Опорное усилие')
    doc.add_paragraph().add_run('F_1=F_2=G/2').font.math = True
    doc.add_paragraph().add_run(f'F_1=F_2={data_in.G}/2={data_out.F1:.2f} H').font.math = True

    doc.add_paragraph('Изгибающий момент над опорами')
    doc.add_paragraph().add_run('M_1=M_2=(q∙e^2)/2-M_0').font.math = True
    doc.add_paragraph().add_run(f'M_1=M_2=({data_out.q:.2f}∙{data_in.e:.2f}^2)/2-{data_out.M0:.2f}={data_out.M1:.2f} Н∙мм').font.math = True

    doc.add_paragraph('Максимальный изгибающий момент между опорами')
    doc.add_paragraph().add_run('M_12=M_0+F_1∙(L/2-a)-q/2∙(L/2+2/3∙H)^2').font.math = True 
    doc.add_paragraph().add_run(f'M_12={data_out.M0:.2f}+{data_out.F1:.2f}∙({data_in.L}/2-{data_in.a})-{data_out.q:.2f}/2∙({data_in.L}/2+2/3∙{data_in.H})^2={data_out.M12:.2f} Н∙мм').font.math = True

    doc.add_paragraph('Поперечное усилие в сечении оболочки над опорой')
    doc.add_paragraph().add_run('Q_1=Q_2=(L-2∙a)/(L+4/3∙H)∙F_1').font.math = True
    doc.add_paragraph().add_run(f'Q_1=Q_2=({data_in.L}-2∙{data_in.a})/({data_in.L}+4/3∙{data_in.H})∙{data_out.F1:.2f}={data_out.Q1:.2f} H').font.math = True

    doc.add_paragraph('Несущую способность обечайки в сечении между опорами следует проверять при условии')
    doc.add_paragraph().add_run('max{M_12}>max{M_1}').font.math = True
    doc.add_paragraph().add_run(f'{data_out.M12:.2f} Н∙мм > {data_out.M1:.2f} Н∙мм').font.math = True
    if data_out.M12 > data_out.M1:

        doc.add_paragraph('Проверка несущей способности обечайки в сечении между опорами')
        doc.add_paragraph('Условие прочности')
        doc.add_paragraph().add_run('(p∙D)/(4∙(s-c))+(4∙M_12∙K_9)/(π∙D^2∙(s-c))≤[σ]∙φ').font.math = True 
        p = doc.add_paragraph('где ')
        p.add_run('K_9').font.math = True
        p.add_run(' - коэффициент, учитывающий частичное заполнение жидкостью')
        doc.add_paragraph().add_run('K_9=max{[1.6-0.20924∙(x-1)+0.028702∙x∙(x-1)+0.4795∙10^3∙y∙(x-1)-0.2391∙10^-6∙x∙y∙(x-1)-0.29936∙10^-2∙(x-1)∙x^2-0.85692∙10^-6∙(х-1)∙у^2+0.88174∙10^-6∙х^2∙(х-1)∙у-0.75955∙10^-8∙у^2∙(х-1)∙х+0.82748∙10^-4∙(х-1)∙х^3+0.48168∙10^-9∙(х-1)∙у^3];1}').font.math = True 

        p = doc.add_paragraph('где ')
        p.add_run('y=D/(s-c);x=L/D').font.math = True
        doc.add_paragraph().add_run(f'y={data_in.D}/({data_in.s}-{data_in.c})={data_out.y:.2f}').font.math = True
        doc.add_paragraph().add_run(f'x={data_in.L}/{data_in.D}={data_out.x:.2f}').font.math = True

        doc.add_paragraph().add_run(f'K_9=max({data_out.K9:.2f};1)={data_out.K9:.2f}').font.math = True

        doc.add_paragraph().add_run(f'(p∙D)/(4∙(s-c))+(4∙M_12∙K_9)/(π∙D^2∙(s-c))=({data_in.p}∙{data_in.D})/(4∙({data_in.s}-{data_in.c}))+(4∙{data_out.M12:.2f}∙{data_out.K9:.2f})/(π∙{data_in.D}^2∙({data_in.s}-{data_in.c}))={data_out.yslproch1_1:.2f}').font.math = True
        doc.add_paragraph().add_run(f'[σ]∙φ={data_in.sigma_d}∙{data_in.fi}={data_out.yslproch1_2:.2f}').font.math = True
        doc.add_paragraph().add_run(f'{data_out.yslproch1_1:.2f}≤{data_out.yslproch1_2:.2f}').font.math = True
        if data_out.yslproch1_1 <= data_out.yslproch1_2:
            doc.add_paragraph('Условие прочности выполняется')
        else:
            doc.add_paragraph().add_run('Условие прочности не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)

        doc.add_paragraph('Условие устойчивости')
        doc.add_paragraph().add_run('|M_12|/[M]≤1').font.math = True

        doc.add_paragraph('где [M] - допускаемый изгибающий момент из условия устойчивости')
        doc.add_paragraph().add_run('[M]=(8.9∙10^-5∙E)/n_y∙D^3∙[(100∙(s-c))/D]^2.5').font.math = True

        doc.add_paragraph().add_run(f'[M]=(8.9∙10^-5∙{data_in.E})/{data_in.ny}∙{data_in.D}^3∙[(100∙({data_in.s}-{data_in.c}))/{data_in.D}]^2.5={data_out.M_d:.2f} Н∙мм').font.math = True
        doc.add_paragraph().add_run(f'|{data_out.M12:.2f}|/{data_out.M_d:.2f}={data_out.yslystoich1:.2f}≤1').font.math = True

        if data_out.yslystoich1 <= 1:
            doc.add_paragraph('Условие устойчивости выполняется')
        else:
            doc.add_paragraph().add_run('Условие устойчивости не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)

    else:
        doc.add_paragraph('Проверка несущей способности обечайки в сечении между опорами не требуется')

    if data_in.type == 1:
        doc.add_paragraph('Проверка несущей способности обечайки, не укрепленной кольцами жесткости в области опорного узла и без подкладного листа в месте опоры')
        doc.add_paragraph('Вспомогательные параметры и коэффициенты')
        doc.add_paragraph('Параметр, определяемый расстоянием от середины опоры до днища')
        doc.add_paragraph().add_run('γ=2.83∙a/D∙√((s-c)/D)').font.math = True
        doc.add_paragraph().add_run(f'γ={data_out.gamma:.2f}').font.math = True

        doc.add_paragraph('Параметр, определяемый шириной пояса опоры')
        doc.add_paragraph().add_run('β_1=0.91∙b/√(D∙(s-c))').font.math = True
        doc.add_paragraph().add_run(f'β_1={data_out.beta1:.2f}').font.math = True

        doc.add_paragraph('Коэффициенты, учитывающие влияние ширины пояса опоры')
        doc.add_paragraph().add_run('K_10=max{(exp(-β_1)∙sin(β_1))/β_1;0.25}').font.math = True
        doc.add_paragraph().add_run(f'K_10=max({data_out.K10_1:.2f};0.25)={data_out.K10:.2f}').font.math = True

        doc.add_paragraph().add_run('K_11=(1-exp(-β_1)∙cos(β_1))/β_1').font.math = True
        doc.add_paragraph().add_run(f'K_11={data_out.K11:.2f}').font.math = True

        doc.add_paragraph('Коэффициенты, учитывающие влияние угла охвата')
        doc.add_paragraph().add_run('K_12=(1.15-0.1432∙δ_1)/sin(0.5∙δ_1)').font.math = True
        doc.add_paragraph().add_run(f'K_12={data_out.K12:.2f}').font.math = True

        doc.add_paragraph().add_run('K_13=(max{1.7-(2.1∙δ_1)/π;0})/sin(0.5∙δ_1)').font.math = True
        doc.add_paragraph().add_run(f'K_13={data_out.K13:.2f}').font.math = True

        doc.add_paragraph().add_run('K_14=(1.45-0.43∙δ_1)/sin(0.5∙δ_1)').font.math = True
        doc.add_paragraph().add_run(f'K_14={data_out.K14:.2f}').font.math = True

        doc.add_paragraph('Коэффициенты, учитывающие влияние расстояния от середины опоры до днища и угла охвата')
        doc.add_paragraph().add_run('K_15=min{1.0;(0.8∙√γ+6∙γ)/δ_1}').font.math = True
        doc.add_paragraph().add_run(f'K_15=min(1.0;{data_out.K15_2:.2f})={data_out.K15:.2f}').font.math = True

        doc.add_paragraph().add_run('K_16=1-0.65/(1+(6∙γ)^2)∙√(π/(3∙δ_1))').font.math = True
        doc.add_paragraph().add_run(f'K_16={data_out.K16:.2f}').font.math = True

        doc.add_paragraph('Коэффициенты, учитывающие влияние ширины пояса опоры и угла охвата')
        doc.add_paragraph().add_run('K_17=1/(1+0.6∙∛(D/(s-c))∙(b/D)∙δ_1)').font.math = True
        doc.add_paragraph().add_run(f'K_17={data_out.K17:.2f}').font.math = True

        doc.add_paragraph('Общее мембранное меридиональное напряжение изгиба от весовых нагрузок, действующее в области опорного узла')
        doc.add_paragraph().add_run('σ ̅_mx=4∙M_i/(π∙D^2∙(s-c))').font.math = True
        doc.add_paragraph().add_run(f'σ ̅_mx={data_out.sigma_mx:.2f}').font.math = True

        doc.add_paragraph('Условие прочности')
        doc.add_paragraph().add_run('F_1≤min{[F]_2;[F]_3}').font.math = True 
        p = doc.add_paragraph('где ')
        p.add_run('[F]_2').font.math = True
        p.add_run(' - допускаемое опорное усилие от нагружения в меридиональном направлении')
        doc.add_paragraph().add_run('[F]_2=(0.7∙[σ_i]_2∙(s-c)∙√(D∙(s-c)))/(K_10∙K_12)').font.math = True 

        p = doc.add_paragraph('   ')
        p.add_run('[F]_3').font.math = True
        p.add_run(' - допускаемое опорное усилие от нагружения в окружном направлении')
        doc.add_paragraph().add_run('[F]_3=(0.9∙[σ_i]_3∙(s-c)∙√(D∙(s-c)))/(K_14∙K_16∙K_17)').font.math = True 

        p = doc.add_paragraph('где ')
        p.add_run('[σ_i]_2, [σ_i]_2').font.math = True
        p.add_run(' - предельные напряжения изгиба в меридиональном и окружном направлениях')

        doc.add_paragraph().add_run('[σ_i]=K_1∙K_2∙[σ]').font.math = True

        doc.add_paragraph().add_run('K_1=(1-ϑ_2^2)/((1/3+ϑ_1∙ϑ_2)+√((1/3+ϑ_1∙ϑ_2)^2+(1-ϑ_2^2)∙ϑ_1^2))').font.math = True
    
        #if 
        p = doc.add_paragraph('')
        p.add_run('K_2=1.25').font.math = True
        p.add_run(' - для рабочих условий')

        p = doc.add_paragraph('для ')
        p.add_run('[σ_i]_2').font.math = True
        doc.add_paragraph().add_run('ϑ_1=-(0,23∙K_13∙K_15)/(K_12∙K_10)').font.math = True
        doc.add_paragraph().add_run(f'ϑ_1={data_out.v1_2:.2f}').font.math = True

        doc.add_paragraph().add_run('ϑ_(2,1)=- ̅σ_mx∙1/(K_2∙[σ])').font.math = True
        doc.add_paragraph().add_run(f'ϑ_(2,1)={data_out.v21_2:.2f}').font.math = True
        doc.add_paragraph().add_run('ϑ_(2,2)=[(p∙D)/(4∙(s-c))- ̅σ_mx]∙1/(K_2∙[σ])').font.math = True
        doc.add_paragraph().add_run(f'ϑ_(2,2)={data_out.v22_2:.2f}').font.math = True

        p = doc.add_paragraph('Для ')
        p.add_run('ϑ_2').font.math = True
        p.add_run('принимают одно из значений ')
        p.add_run('ϑ_(2,1)').font.math = True
        p.add_run(' или ')
        p.add_run('ϑ_(2,2)').font.math = True
        p.add_run(', для которого предельное напряжение изгибабудет наименьшим.')
        if data_out.K1_21 < data_out.K1_22:
            doc.add_paragraph().add_run(f'ϑ_2=ϑ_(2,1)={data_out.v21_2:.2f}').font.math = True
        else:
            doc.add_paragraph().add_run(f'ϑ_2=ϑ_(2,2)={data_out.v22_2:.2f}').font.math = True
        
        doc.add_paragraph().add_run(f'K_1={data_out.K1_2:.2f}').font.math = True

        doc.add_paragraph().add_run(f'[σ_i]_2={data_out.K1_2:.2f}∙{data_out.K2:.2f}∙{data_in.sigma_d}={data_out.sigmai2:.2f}').font.math = True

        doc.add_paragraph().add_run(f'[F]_2=(0.7∙{data_out.sigmai2:.2f}∙({data_in.s}-{data_in.c})∙√({data_in.D}∙({data_in.s}-{data_in.c})))/({data_out.K10:.2f}∙{data_out.K12:.2f})={data_out.F_d2:.2f}').font.math = True


        p = doc.add_paragraph('для ')
        p.add_run('[σ_i]_3').font.math = True
        doc.add_paragraph().add_run('ϑ_1=-(0,53∙K_11)/(K_14∙K_16∙K_17∙sin(0.5∙δ_1))').font.math = True
        doc.add_paragraph().add_run(f'ϑ_1={data_out.v1_3:.2f}').font.math = True

        doc.add_paragraph().add_run('ϑ_(2,1)=0').font.math = True
        doc.add_paragraph().add_run(f'ϑ_(2,1)={data_out.v21_3:.2f}').font.math = True
        doc.add_paragraph().add_run('ϑ_(2,2)=(p∙D)/(2∙(s-c))∙1/(K_2∙[σ])').font.math = True
        doc.add_paragraph().add_run(f'ϑ_(2,2)={data_out.v22_3:.2f}').font.math = True

        p = doc.add_paragraph('Для ')
        p.add_run('ϑ_2').font.math = True
        p.add_run('принимают одно из значений ')
        p.add_run('ϑ_(2,1)').font.math = True
        p.add_run(' или ')
        p.add_run('ϑ_(2,2)').font.math = True
        p.add_run(', для которого предельное напряжение изгибабудет наименьшим.')
        if data_out.K1_31 < data_out.K1_32:
            doc.add_paragraph().add_run(f'ϑ_2=ϑ_(2,1)={data_out.v21_3:.2f}').font.math = True
        else:
            doc.add_paragraph().add_run(f'ϑ_2=ϑ_(2,2)={data_out.v22_3:.2f}').font.math = True
        
        doc.add_paragraph().add_run(f'K_1={data_out.K1_3:.2f}').font.math = True

        doc.add_paragraph().add_run(f'[σ_i]_3={data_out.K1_3:.2f}∙{data_out.K2:.2f}∙{data_in.sigma_d}={data_out.sigmai3:.2f}').font.math = True

        doc.add_paragraph().add_run(f'[F]_3=(0.9∙{data_out.sigmai2:.2f}∙({data_in.s}-{data_in.c})∙√({data_in.D}∙({data_in.s}-{data_in.c})))/({data_out.K14:.2f}∙{data_out.K16:.2f}∙{data_out.K17:.2f})={data_out.F_d3:.2f}').font.math = True

        doc.add_paragraph().add_run(f'{data_out.F1:.2f}≤min({data_out.F_d2:.2f};{data_out.F_d3:.2f})').font.math = True
        if data_out.F1 <= min(data_out.F_d2, data_out.F_d3):
            doc.add_paragraph('Условие прочности выполняется')
        else:
            doc.add_paragraph().add_run('Условие прочности не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)

        doc.add_paragraph('Условие устойчивости')
        #добавить наружное давление
        doc.add_paragraph().add_run('|p|/[p]+|M_i|/[M]+|F_e|/[F]+(Q_i/[Q])^2≤1').font.math = True

        p = doc.add_paragraph('где ')
        p.add_run('F_e').font.math = True
        p.add_run(' - эффективное осевое усилие от местных мембранных напряжений, действующих в области опоры')
        doc.add_paragraph().add_run('F_e=F_i∙π/4∙K_13∙K_15∙√(D/(s-c))').font.math = True 
        doc.add_paragraph().add_run(f'F_e={data_out.F1:.2f}∙π/4∙{data_out.K13:.2f}∙{data_out.K15:.2f}∙√({data_in.D}/({data_in.s}-{data_in.c}))={data_out.Fe:.2f}').font.math = True
        doc.add_paragraph().add_run(f'{data_out.M1:.2f}/{data_out.M_d:.2f}+{data_out.Fe:.2f}/{data_out.F_d:.2f}+({data_out.Q1:.2f}/{data_out.Q_d:.2f})^2={data_out.yslystoich2:.2f}≤1').font.math = True

        if data_out.yslystoich2 <= 1:
            doc.add_paragraph('Условие устойчивости выполняется')
        else:
            doc.add_paragraph().add_run('Условие устойчивости не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)

    elif data_in.type == 2:
        doc.add_paragraph('Проверка несущей способности обечайки, не укрепленной кольцами жесткости в области опорного узла с подкладным листом в месте опоры')
        doc.add_paragraph('Вспомогательные параметры и коэффициенты')

        doc.add_paragraph().add_run('s_ef=(s-c)∙√(1+(s_2/(s-c))^2)').font.math = True
        doc.add_paragraph().add_run(f's_ef={data_out.sef:.2f}').font.math = True

        doc.add_paragraph('Параметр, определяемый расстоянием от середины опоры до днища')
        doc.add_paragraph().add_run('γ=2.83∙a/D∙√(s_ef)/D').font.math = True
        doc.add_paragraph().add_run(f'γ={data_out.gamma:.2f}').font.math = True

        doc.add_paragraph('Параметр, определяемый шириной пояса опоры')
        doc.add_paragraph().add_run('β_1=0.91∙b_2/√(D∙(s_ef))').font.math = True
        doc.add_paragraph().add_run(f'β_1={data_out.beta1:.2f}').font.math = True

        doc.add_paragraph('Коэффициенты, учитывающие влияние ширины пояса опоры')
        doc.add_paragraph().add_run('K_10=max{(exp(-β_1)∙sin(β_1))/β_1;0.25}').font.math = True
        doc.add_paragraph().add_run(f'K_10=max({data_out.K10_1:.2f};0.25)={data_out.K10:.2f}').font.math = True

        doc.add_paragraph().add_run('K_11=(1-exp(-β_1)∙cos(β_1))/β_1').font.math = True
        doc.add_paragraph().add_run(f'K_11={data_out.K11:.2f}').font.math = True

        doc.add_paragraph('Коэффициенты, учитывающие влияние угла охвата')
        doc.add_paragraph().add_run('K_12=(1.15-0.1432∙δ_1)/sin(0.5∙δ_1)').font.math = True
        doc.add_paragraph().add_run(f'K_12={data_out.K12:.2f}').font.math = True

        doc.add_paragraph().add_run('K_13=(max{1.7-(2.1∙δ_1)/π;0})/sin(0.5∙δ_1)').font.math = True
        doc.add_paragraph().add_run(f'K_13={data_out.K13:.2f}').font.math = True

        doc.add_paragraph().add_run('K_14=(1.45-0.43∙δ_1)/sin(0.5∙δ_1)').font.math = True
        doc.add_paragraph().add_run(f'K_14={data_out.K14:.2f}').font.math = True

        doc.add_paragraph('Коэффициенты, учитывающие влияние расстояния от середины опоры до днища и угла охвата')
        doc.add_paragraph().add_run('K_15=min{1.0;(0.8∙√γ+6∙γ)/δ_1}').font.math = True
        doc.add_paragraph().add_run(f'K_15=min(1.0;{data_out.K15_2:.2f})={data_out.K15:.2f}').font.math = True

        doc.add_paragraph().add_run('K_16=1-0.65/(1+(6∙γ)^2)∙√(π/(3∙δ_1))').font.math = True
        doc.add_paragraph().add_run(f'K_16={data_out.K16:.2f}').font.math = True

        doc.add_paragraph('Коэффициенты, учитывающие влияние ширины пояса опоры и угла охвата')
        doc.add_paragraph().add_run('K_17=1/(1+0.6∙∛(D/(s_ef))∙(b_2/D)∙δ_1)').font.math = True
        doc.add_paragraph().add_run(f'K_17={data_out.K17:.2f}').font.math = True

        doc.add_paragraph('Общее мембранное меридиональное напряжение изгиба от весовых нагрузок, действующее в области опорного узла')
        doc.add_paragraph().add_run('σ ̅_mx=4∙M_i/(π∙D^2∙(s_ef))').font.math = True
        doc.add_paragraph().add_run(f'σ ̅_mx={data_out.sigma_mx:.2f}').font.math = True

        doc.add_paragraph('Условие прочности')
        doc.add_paragraph().add_run('F_1≤min{[F]_2;[F]_3}').font.math = True 
        p = doc.add_paragraph('где ')
        p.add_run('[F]_2').font.math = True
        p.add_run(' - допускаемое опорное усилие от нагружения в меридиональном направлении')
        doc.add_paragraph().add_run('[F]_2=(0.7∙[σ_i]_2∙(s_ef)∙√(D∙(s_ef)))/(K_10∙K_12)').font.math = True 

        p = doc.add_paragraph('   ')
        p.add_run('[F]_3').font.math = True
        p.add_run(' - допускаемое опорное усилие от нагружения в окружном направлении')
        doc.add_paragraph().add_run('[F]_3=(0.9∙[σ_i]_3∙(s_ef)∙√(D∙(s_ef)))/(K_14∙K_16∙K_17)').font.math = True 

        p = doc.add_paragraph('где ')
        p.add_run('[σ_i]_2, [σ_i]_2').font.math = True
        p.add_run(' - предельные напряжения изгиба в меридиональном и окружном направлениях')

        doc.add_paragraph().add_run('[σ_i]=K_1∙K_2∙[σ]').font.math = True

        doc.add_paragraph().add_run('K_1=(1-ϑ_2^2)/((1/3+ϑ_1∙ϑ_2)+√((1/3+ϑ_1∙ϑ_2)^2+(1-ϑ_2^2)∙ϑ_1^2))').font.math = True
    
        #if 
        p = doc.add_paragraph('')
        p.add_run('K_2=1.25').font.math = True
        p.add_run(' - для рабочих условий')

        p = doc.add_paragraph('для ')
        p.add_run('[σ_i]_2').font.math = True
        doc.add_paragraph().add_run('ϑ_1=-(0,23∙K_13∙K_15)/(K_12∙K_10)').font.math = True
        doc.add_paragraph().add_run(f'ϑ_1={data_out.v1_2:.2f}').font.math = True

        doc.add_paragraph().add_run('ϑ_(2,1)=- ̅σ_mx∙1/(K_2∙[σ])').font.math = True
        doc.add_paragraph().add_run(f'ϑ_(2,1)={data_out.v21_2:.2f}').font.math = True
        doc.add_paragraph().add_run('ϑ_(2,2)=[(p∙D)/(4∙(s_ef))- ̅σ_mx]∙1/(K_2∙[σ])').font.math = True
        doc.add_paragraph().add_run(f'ϑ_(2,2)={data_out.v22_2:.2f}').font.math = True

        p = doc.add_paragraph('Для ')
        p.add_run('ϑ_2').font.math = True
        p.add_run('принимают одно из значений ')
        p.add_run('ϑ_(2,1)').font.math = True
        p.add_run(' или ')
        p.add_run('ϑ_(2,2)').font.math = True
        p.add_run(', для которого предельное напряжение изгибабудет наименьшим.')
        if data_out.K1_21 < data_out.K1_22:
            doc.add_paragraph().add_run(f'ϑ_2=ϑ_(2,1)={data_out.v21_2:.2f}').font.math = True
        else:
            doc.add_paragraph().add_run(f'ϑ_2=ϑ_(2,2)={data_out.v22_2:.2f}').font.math = True
        
        doc.add_paragraph().add_run(f'K_1={data_out.K1_2:.2f}').font.math = True

        doc.add_paragraph().add_run(f'[σ_i]_2={data_out.K1_2:.2f}∙{data_out.K2:.2f}∙{data_in.sigma_d}={data_out.sigmai2:.2f}').font.math = True

        doc.add_paragraph().add_run(f'[F]_2=(0.7∙{data_out.sigmai2:.2f}∙({data_out.sef:.2f})∙√({data_in.D}∙({data_out.sef:.2f})))/({data_out.K10:.2f}∙{data_out.K12:.2f})={data_out.F_d2:.2f}').font.math = True


        p = doc.add_paragraph('для ')
        p.add_run('[σ_i]_3').font.math = True
        doc.add_paragraph().add_run('ϑ_1=-(0,53∙K_11)/(K_14∙K_16∙K_17∙sin(0.5∙δ_1))').font.math = True
        doc.add_paragraph().add_run(f'ϑ_1={data_out.v1_3:.2f}').font.math = True

        doc.add_paragraph().add_run('ϑ_(2,1)=0').font.math = True
        doc.add_paragraph().add_run(f'ϑ_(2,1)={data_out.v21_3:.2f}').font.math = True
        doc.add_paragraph().add_run('ϑ_(2,2)=(p∙D)/(2∙(s_ef))∙1/(K_2∙[σ])').font.math = True
        doc.add_paragraph().add_run(f'ϑ_(2,2)={data_out.v22_3:.2f}').font.math = True

        p = doc.add_paragraph('Для ')
        p.add_run('ϑ_2').font.math = True
        p.add_run('принимают одно из значений ')
        p.add_run('ϑ_(2,1)').font.math = True
        p.add_run(' или ')
        p.add_run('ϑ_(2,2)').font.math = True
        p.add_run(', для которого предельное напряжение изгибабудет наименьшим.')
        if data_out.K1_31 < data_out.K1_32:
            doc.add_paragraph().add_run(f'ϑ_2=ϑ_(2,1)={data_out.v21_3:.2f}').font.math = True
        else:
            doc.add_paragraph().add_run(f'ϑ_2=ϑ_(2,2)={data_out.v22_3:.2f}').font.math = True
        
        doc.add_paragraph().add_run(f'K_1={data_out.K1_3:.2f}').font.math = True

        doc.add_paragraph().add_run(f'[σ_i]_3={data_out.K1_3:.2f}∙{data_out.K2:.2f}∙{data_in.sigma_d}={data_out.sigmai3:.2f}').font.math = True

        doc.add_paragraph().add_run(f'[F]_3=(0.9∙{data_out.sigmai2:.2f}∙({data_out.sef:.2f})∙√({data_in.D}∙({data_out.sef:.2f})))/({data_out.K14:.2f}∙{data_out.K16:.2f}∙{data_out.K17:.2f})={data_out.F_d3:.2f}').font.math = True

        doc.add_paragraph().add_run(f'{data_out.F1:.2f}≤min({data_out.F_d2:.2f};{data_out.F_d3:.2f})').font.math = True
        if data_out.F1 <= min(data_out.F_d2, data_out.F_d3):
            doc.add_paragraph('Условие прочности выполняется')
        else:
            doc.add_paragraph().add_run('Условие прочности не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)

        doc.add_paragraph('Условие устойчивости')
        #добавить наружное давление
        doc.add_paragraph().add_run('|p|/[p]+|M_i|/[M]+|F_e|/[F]+(Q_i/[Q])^2≤1').font.math = True

        p = doc.add_paragraph('где ')
        p.add_run('F_e').font.math = True
        p.add_run(' - эффективное осевое усилие от местных мембранных напряжений, действующих в области опоры')
        doc.add_paragraph().add_run('F_e=F_i∙π/4∙K_13∙K_15∙√(D/(s_ef))').font.math = True 
        doc.add_paragraph().add_run(f'F_e={data_out.F1:.2f}∙π/4∙{data_out.K13:.2f}∙{data_out.K15:.2f}∙√({data_in.D}/({data_out.sef:.2f}))={data_out.Fe:.2f}').font.math = True
        doc.add_paragraph().add_run(f'{data_out.M1:.2f}/{data_out.M_d:.2f}+{data_out.Fe:.2f}/{data_out.F_d:.2f}+({data_out.Q1:.2f}/{data_out.Q_d:.2f})^2={data_out.yslystoich2:.2f}≤1').font.math = True

        if data_out.yslystoich2 <= 1:
            doc.add_paragraph('Условие устойчивости выполняется')
        else:
            doc.add_paragraph().add_run('Условие устойчивости не выполняется').font.color.rgb = docx.shared.RGBColor(255,0,0)
    

    doc.save(docum)